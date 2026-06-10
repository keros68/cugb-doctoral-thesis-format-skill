#!/usr/bin/env python
"""Generate a one-click local precheck report for CUGB doctoral thesis DOCX files."""

from __future__ import annotations

import argparse
import collections
import datetime as dt
import json
import platform
import re
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import check_cugb_doctoral_docx as checker
import extract_cugb_docx_format as extractor
import summarize_cugb_detection_report as report_summarizer


RESIDUE_PATTERNS = [
    "中国民航",
    "中南大学",
    "Central South",
    "PLACEHOLDER",
    "TODO",
]


def now_text() -> str:
    return dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def finding_to_dict(finding) -> dict:
    return {"level": finding.level, "location": finding.location, "message": finding.message}


def count_findings(findings) -> dict:
    counts = collections.Counter(finding.level for finding in findings)
    return {level: counts.get(level, 0) for level in ("ERROR", "WARN", "REMIND", "INFO")}


def status_from_counts(counts: dict) -> str:
    if counts.get("ERROR", 0):
        return "must-fix"
    if counts.get("WARN", 0):
        return "review-needed"
    if counts.get("REMIND", 0):
        return "polish-needed"
    return "clean-local"


def paragraph_count(doc: Document) -> int:
    return sum(1 for paragraph in doc.paragraphs if paragraph.text.strip())


def section_summary(doc: Document) -> list[dict]:
    return extractor.section_report(doc)


def hidden_header_residue(docx_path: Path) -> list[dict]:
    hits = []
    with zipfile.ZipFile(docx_path) as package:
        for name in package.namelist():
            if not re.match(r"word/(header|footer)\d*\.xml$", name):
                continue
            data = package.read(name).decode("utf-8", errors="ignore")
            text = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", data))
            text = clean(text)
            for pattern in RESIDUE_PATTERNS:
                if pattern in text:
                    hits.append({"part": name, "pattern": pattern, "text": text[:200]})
    return hits


def word_header_audit(docx_path: Path) -> dict:
    if platform.system() != "Windows":
        return {"available": False, "reason": "Word COM audit only runs on Windows"}

    ps_script = r"""
param([string]$Path)
$ErrorActionPreference = 'Stop'
[Console]::OutputEncoding = [System.Text.Encoding]::UTF8
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
try {
  $doc = $word.Documents.Open($Path, $false, $true)
  $sections = @()
  for ($i=1; $i -le $doc.Sections.Count; $i++) {
    $sec = $doc.Sections.Item($i)
    $odd = ($sec.Headers.Item(1).Range.Text -replace '[\r\a\f]', ' ' -replace '\s+', ' ').Trim()
    $even = ($sec.Headers.Item(3).Range.Text -replace '[\r\a\f]', ' ' -replace '\s+', ' ').Trim()
    $first = ($sec.Headers.Item(2).Range.Text -replace '[\r\a\f]', ' ' -replace '\s+', ' ').Trim()
    $sections += [PSCustomObject]@{
      section = $i
      odd_header = $odd
      even_header = $even
      first_header = $first
      footer_distance_cm = [Math]::Round($sec.PageSetup.FooterDistance / 28.3464567, 2)
      page_width_cm = [Math]::Round($sec.PageSetup.PageWidth / 28.3464567, 2)
      page_height_cm = [Math]::Round($sec.PageSetup.PageHeight / 28.3464567, 2)
    }
  }
  $doc.Close($false)
  [PSCustomObject]@{ available = $true; sections = $sections } | ConvertTo-Json -Depth 6 -Compress
}
finally {
  $word.Quit()
}
"""
    with tempfile.NamedTemporaryFile("w", suffix=".ps1", delete=False, encoding="utf-8") as handle:
        handle.write(ps_script)
        script_path = Path(handle.name)

    try:
        result = subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(script_path),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=180,
        )
    finally:
        script_path.unlink(missing_ok=True)

    if result.returncode != 0:
        return {
            "available": False,
            "reason": "Word COM audit failed",
            "stderr": result.stderr.strip()[-1000:],
        }
    try:
        return json.loads(result.stdout)
    except json.JSONDecodeError:
        return {
            "available": False,
            "reason": "Word COM audit returned non-JSON output",
            "stdout": result.stdout.strip()[-1000:],
        }


def word_header_issues(audit: dict) -> list[str]:
    if not audit.get("available"):
        return []
    issues = []
    for section in audit.get("sections", []):
        number = section.get("section")
        odd = section.get("odd_header", "")
        even = section.get("even_header", "")
        footer_distance = section.get("footer_distance_cm")
        if number and number >= 3 and odd and "中国地质大学（北京）博士学位论文" not in odd:
            issues.append(f"第 {number} 节奇数页页眉异常：{odd}")
        if number and number >= 3 and even and any(pattern in even for pattern in RESIDUE_PATTERNS):
            issues.append(f"第 {number} 节偶数页页眉疑似残留：{even}")
        if footer_distance is not None and abs(float(footer_distance) - 1.5) > 0.08:
            issues.append(f"第 {number} 节页脚距边界为 {footer_distance} cm，模板约 1.50 cm")
    return issues


def filter_findings_for_report(findings, word_audit: dict, word_issues: list[str]):
    if not word_audit.get("available") or word_issues:
        return list(findings)
    filtered = []
    for finding in findings:
        if finding.level == "WARN" and "python-docx 视图中重复" in finding.message:
            continue
        filtered.append(finding)
    return filtered


def top_findings(findings, level: str, limit: int = 12) -> list[dict]:
    rows = [finding_to_dict(finding) for finding in findings if finding.level == level]
    return rows[:limit]


def manual_review_items(result: dict) -> list[str]:
    items = [
        "在 Word 中全选并更新域，重点更新目录页码。",
        "导出 PDF 后抽查封面、声明、摘要、目录、第 1 章首页、每章首页、图表密集页、参考文献、致谢、个人简历。",
        "学校检测机会有限时，先把本报告中的 ERROR 清零，再处理 WARN 和高频 REMIND。",
    ]
    if any(section.get("page_width_cm") == 29.7 for section in result["structure"]["sections"]):
        items.append("文档含横向 A4 section，多半用于宽表；提交前人工确认横向页在学校系统中可接受。")
    if result.get("word_header_issues"):
        items.append("Word COM 页眉审计仍有异常，需先修 section 页眉页脚。")
    return items


def build_markdown(result: dict) -> str:
    counts = result["finding_counts"]
    lines = [
        "# CUGB 博士论文本地预检报告",
        "",
        f"- 文件：`{result['input_file']}`",
        f"- 生成时间：{result['generated_at']}",
        f"- 本地状态：{result['status_label']}",
        f"- 检查统计：ERROR {counts['ERROR']}，WARN {counts['WARN']}，REMIND {counts['REMIND']}，INFO {counts['INFO']}",
        f"- 结构概况：{result['structure']['section_count']} 个 section，{result['structure']['paragraph_count']} 个非空段落，{result['structure']['table_count']} 个表格",
        f"- Word 域：TOC {result['structure']['xml_counts']['toc_fields']}，PAGE {result['structure']['xml_counts']['page_fields']}，OMath {result['structure']['xml_counts']['omath_objects']}",
        "",
        "## 必须修",
    ]
    if result["top_errors"]:
        for item in result["top_errors"]:
            lines.append(f"- {item['location']}：{item['message']}")
    else:
        lines.append("- 未发现脚本级 ERROR。")

    lines.extend(["", "## 建议修"])
    if result["top_warnings"]:
        for item in result["top_warnings"]:
            lines.append(f"- {item['location']}：{item['message']}")
    else:
        lines.append("- 未发现脚本级 WARN。")

    lines.extend(["", "## 学校系统可能提醒"])
    if result["top_reminders"]:
        for item in result["top_reminders"]:
            lines.append(f"- {item['location']}：{item['message']}")
    else:
        lines.append("- 未发现脚本级 REMIND。")

    lines.extend(["", "## Word 真实页眉页脚"])
    audit = result["word_header_audit"]
    if audit.get("available"):
        if result["word_header_issues"]:
            for issue in result["word_header_issues"]:
                lines.append(f"- {issue}")
        else:
            lines.append("- Word COM 未发现页眉页脚硬风险。")
        for section in audit.get("sections", [])[:20]:
            lines.append(
                f"- 第 {section['section']} 节：奇数页 `{section.get('odd_header','')}`；偶数页 `{section.get('even_header','')}`；页脚距 {section.get('footer_distance_cm')} cm"
            )
    else:
        lines.append(f"- 未运行或不可用：{audit.get('reason', 'skipped')}")

    lines.extend(["", "## 隐藏页眉/页脚残留"])
    if result["hidden_header_residue"]:
        for hit in result["hidden_header_residue"]:
            lines.append(f"- {hit['part']} 命中 `{hit['pattern']}`：{hit['text']}")
    else:
        lines.append("- 未发现外校/TODO 等隐藏页眉页脚残留。")

    if result.get("school_detection_report"):
        school = result["school_detection_report"]
        lines.extend(
            [
                "",
                "## 学校检测报告汇总",
                f"- 总条数：{school['total_rows']}",
                f"- 等级：{school['levels']}",
                f"- 高频类别：{school['categories']}",
            ]
        )

    lines.extend(["", "## 人工复核清单"])
    for item in result["manual_review"]:
        lines.append(f"- {item}")

    return "\n".join(lines) + "\n"


def build_precheck_report(
    docx_path: Path | str,
    output_dir: Path | str | None = None,
    detection_report: Path | str | None = None,
    include_word: bool = True,
) -> dict:
    docx_path = Path(docx_path)
    if output_dir is None:
        output_dir = docx_path.parent / "cugb-precheck"
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document(docx_path)
    raw_findings = checker.check_document(docx_path)
    word_audit = word_header_audit(docx_path) if include_word else {"available": False, "reason": "skipped by --no-word"}
    word_issues = word_header_issues(word_audit)
    findings = filter_findings_for_report(raw_findings, word_audit, word_issues)
    counts = count_findings(findings)
    status = status_from_counts(counts)
    status_labels = {
        "must-fix": "必须先修 ERROR",
        "review-needed": "无 ERROR，但需复核 WARN",
        "polish-needed": "仅剩提醒项，适合提交前润色",
        "clean-local": "本地脚本未发现问题",
    }

    result = {
        "input_file": str(docx_path),
        "generated_at": now_text(),
        "status": status,
        "status_label": status_labels[status],
        "finding_counts": counts,
        "findings": [finding_to_dict(finding) for finding in findings],
        "raw_findings": [finding_to_dict(finding) for finding in raw_findings],
        "top_errors": top_findings(findings, "ERROR"),
        "top_warnings": top_findings(findings, "WARN"),
        "top_reminders": top_findings(findings, "REMIND", limit=20),
        "structure": {
            "section_count": len(doc.sections),
            "paragraph_count": paragraph_count(doc),
            "table_count": len(doc.tables),
            "sections": section_summary(doc),
            "xml_counts": extractor.xml_counts(docx_path),
        },
        "word_header_audit": word_audit,
        "word_header_issues": word_issues,
        "hidden_header_residue": hidden_header_residue(docx_path),
        "school_detection_report": None,
        "outputs": {},
    }
    result["manual_review"] = manual_review_items(result)

    if detection_report:
        summary = report_summarizer.parse_report(Path(detection_report))
        summary.pop("rows", None)
        result["school_detection_report"] = summary

    base = f"{docx_path.stem}_cugb-precheck"
    json_path = output_dir / f"{base}.json"
    md_path = output_dir / f"{base}.md"
    result["outputs"] = {"json": str(json_path), "markdown": str(md_path)}

    json_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    md_path.write_text(build_markdown(result), encoding="utf-8")
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output-dir", type=Path)
    parser.add_argument("--detection-report", type=Path, help="Optional school HTML detection report")
    parser.add_argument("--no-word", action="store_true", help="Skip Microsoft Word COM audit")
    args = parser.parse_args()

    result = build_precheck_report(
        args.docx,
        output_dir=args.output_dir,
        detection_report=args.detection_report,
        include_word=not args.no_word,
    )
    print(f"Markdown: {result['outputs']['markdown']}")
    print(f"JSON: {result['outputs']['json']}")
    print(f"Status: {result['status_label']}")
    return 1 if result["finding_counts"].get("ERROR", 0) else 0


if __name__ == "__main__":
    sys.exit(main())
