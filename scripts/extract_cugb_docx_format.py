#!/usr/bin/env python
"""Extract a compact CUGB doctoral thesis DOCX structure report."""

from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document


def cm(value):
    return None if value is None else round(value.cm, 2)


def text_of_paragraphs(paragraphs):
    return " | ".join(" ".join(p.text.split()) for p in paragraphs if p.text.strip())


def section_report(doc: Document):
    rows = []
    for index, section in enumerate(doc.sections, 1):
        rows.append(
            {
                "section": index,
                "page_width_cm": cm(section.page_width),
                "page_height_cm": cm(section.page_height),
                "top_margin_cm": cm(section.top_margin),
                "bottom_margin_cm": cm(section.bottom_margin),
                "left_margin_cm": cm(section.left_margin),
                "right_margin_cm": cm(section.right_margin),
                "header_distance_cm": cm(section.header_distance),
                "footer_distance_cm": cm(section.footer_distance),
                "different_first_page": bool(section.different_first_page_header_footer),
                "header": text_of_paragraphs(section.header.paragraphs),
                "even_header": text_of_paragraphs(section.even_page_header.paragraphs),
                "footer": text_of_paragraphs(section.footer.paragraphs),
            }
        )
    return rows


def paragraph_report(doc: Document, limit: int):
    rows = []
    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        fmt = paragraph.paragraph_format
        runs = []
        for run in paragraph.runs[:6]:
            if not run.text.strip():
                continue
            runs.append(
                {
                    "text": run.text.strip()[:40],
                    "font": run.font.name,
                    "size_pt": None if run.font.size is None else round(run.font.size.pt, 1),
                    "bold": run.font.bold,
                    "superscript": run.font.superscript,
                    "subscript": run.font.subscript,
                }
            )
        rows.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": text[:220],
                "alignment": str(paragraph.alignment or fmt.alignment),
                "first_line_indent_cm": cm(fmt.first_line_indent),
                "line_spacing": fmt.line_spacing,
                "runs": runs,
            }
        )
        if len(rows) >= limit:
            break
    return rows


def table_report(doc: Document, limit: int):
    rows = []
    for index, table in enumerate(doc.tables[:limit], 1):
        preview = []
        for row in table.rows[:8]:
            preview.append([" ".join(cell.text.split())[:80] for cell in row.cells])
        rows.append({"index": index, "rows": len(table.rows), "cols": len(table.columns), "preview": preview})
    return rows


def xml_counts(path: Path):
    counts = {"toc_fields": 0, "page_fields": 0, "omath_objects": 0}
    with zipfile.ZipFile(path) as package:
        for name in package.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            data = package.read(name).decode("utf-8", errors="ignore")
            counts["toc_fields"] += data.count("TOC ")
            counts["page_fields"] += data.count(" PAGE ")
            counts["omath_objects"] += data.count("<m:oMath")
    return counts


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--paragraph-limit", type=int, default=120)
    parser.add_argument("--table-limit", type=int, default=8)
    args = parser.parse_args()

    doc = Document(args.docx)
    report = {
        "file": str(args.docx),
        "sections": section_report(doc),
        "xml_counts": xml_counts(args.docx),
        "paragraphs": paragraph_report(doc, args.paragraph_limit),
        "tables": table_report(doc, args.table_limit),
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
