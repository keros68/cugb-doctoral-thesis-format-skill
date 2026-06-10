from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm


SCRIPTS_DIR = Path(__file__).resolve().parent


def load_module(name: str, filename: str):
    spec = importlib.util.spec_from_file_location(name, SCRIPTS_DIR / filename)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def finding_messages(findings):
    return "\n".join(f"{item.level} {item.location} {item.message}" for item in findings)


def test_checker_flags_cover_line_spacing_and_field_indent(tmp_path):
    checker = load_module("checker", "check_cugb_doctoral_docx.py")
    doc = Document()
    school = doc.add_paragraph("中国地质大学（北京）")
    school.paragraph_format.line_spacing = 1.5
    field = doc.add_paragraph("学号：123")
    field.paragraph_format.first_line_indent = Cm(3.0)
    path = tmp_path / "bad_cover.docx"
    doc.save(path)

    messages = finding_messages(checker.check_document(path))

    assert "封面学校名行应保持单倍行距" in messages
    assert "封面字段首行缩进应约为 2.22 cm" in messages


def test_checker_flags_toc_items_with_heading_font(tmp_path):
    checker = load_module("checker", "check_cugb_doctoral_docx.py")
    doc = Document()
    doc.add_paragraph("目 录")
    toc_item = doc.add_paragraph()
    run = toc_item.add_run("摘要 1")
    run.font.name = "黑体"
    doc.add_paragraph("1 引言")
    path = tmp_path / "bad_toc.docx"
    doc.save(path)

    messages = finding_messages(checker.check_document(path))

    assert "目录项不应使用黑体" in messages


def test_checker_does_not_treat_dates_or_numbered_lists_as_headings(tmp_path):
    checker = load_module("checker", "check_cugb_doctoral_docx.py")
    doc = Document()
    doc.add_paragraph("2023年07月")
    doc.add_paragraph("2010年9月-2014年6月 河南理工大学 本科")
    doc.add_paragraph("1. 第一作者（1篇）")
    path = tmp_path / "dates_and_lists.docx"
    doc.save(path)

    messages = finding_messages(checker.check_document(path))

    assert "标题编号与标题文字之间应有 1 个空格" not in messages


def test_checker_warns_for_landscape_a4_without_page_size_error(tmp_path):
    checker = load_module("checker", "check_cugb_doctoral_docx.py")
    doc = Document()
    section = doc.sections[-1]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    doc.add_paragraph("宽表")
    path = tmp_path / "landscape.docx"
    doc.save(path)

    messages = finding_messages(checker.check_document(path))

    assert "检测到横向 A4 section" in messages
    assert "纸张宽度应约为 21.0 cm" not in messages


def test_detection_report_summary_counts_issue_categories(tmp_path):
    summarizer = load_module("summarizer", "summarize_cugb_detection_report.py")
    html = """
    <html><body><table>
      <tr><th>序号</th><th>等级</th><th>位置</th><th>问题描述</th></tr>
      <tr><td>1</td><td>错误</td><td>第1页</td><td>行距值(要求：单倍行距, 实际：1.5 倍行距)</td></tr>
      <tr><td>2</td><td>严重</td><td>关键词</td><td>关键词{ A }分隔符 { ； }不符合规范要求{ ， }</td></tr>
      <tr><td>3</td><td>提醒</td><td>参考文献</td><td>连续出版物的析出文献应有页码</td></tr>
    </table></body></html>
    """
    path = tmp_path / "report.html"
    path.write_text(html, encoding="utf-8")

    summary = summarizer.parse_report(path)

    assert summary["total_rows"] == 3
    assert summary["levels"]["错误"] == 1
    assert summary["levels"]["严重"] == 1
    assert summary["categories"]["cover-line-spacing"] == 1
    assert summary["categories"]["keywords"] == 1
    assert summary["categories"]["reference-pages"] == 1


def test_precheck_generates_markdown_and_json_report(tmp_path):
    precheck = load_module("precheck", "precheck_cugb_doctoral_thesis.py")
    doc = Document()
    school = doc.add_paragraph("中国地质大学（北京）")
    school.paragraph_format.line_spacing = 1.5
    doc.add_paragraph("中文摘要")
    path = tmp_path / "bad_precheck.docx"
    doc.save(path)

    result = precheck.build_precheck_report(path, output_dir=tmp_path, include_word=False)

    md_path = Path(result["outputs"]["markdown"])
    json_path = Path(result["outputs"]["json"])
    assert md_path.exists()
    assert json_path.exists()
    assert result["status"] == "must-fix"
    assert result["finding_counts"]["ERROR"] >= 1
    assert "CUGB 博士论文本地预检报告" in md_path.read_text(encoding="utf-8")
    assert "必须修" in md_path.read_text(encoding="utf-8")


def test_precheck_includes_detection_report_summary(tmp_path):
    precheck = load_module("precheck", "precheck_cugb_doctoral_thesis.py")
    doc = Document()
    doc.add_paragraph("中国地质大学（北京）")
    docx_path = tmp_path / "sample.docx"
    doc.save(docx_path)
    html = """
    <html><body><table>
      <tr><th>序号</th><th>等级</th><th>位置</th><th>问题描述</th></tr>
      <tr><td>1</td><td>严重</td><td>页眉页脚</td><td>第5节：偶数页页眉内容错误</td></tr>
    </table></body></html>
    """
    report_path = tmp_path / "school-report.html"
    report_path.write_text(html, encoding="utf-8")

    result = precheck.build_precheck_report(
        docx_path,
        output_dir=tmp_path,
        detection_report=report_path,
        include_word=False,
    )

    assert result["school_detection_report"]["total_rows"] == 1
    assert result["school_detection_report"]["levels"]["严重"] == 1


def test_precheck_suppresses_python_docx_header_warning_when_word_audit_is_clean():
    precheck = load_module("precheck", "precheck_cugb_doctoral_thesis.py")
    checker = load_module("checker_for_filter", "check_cugb_doctoral_docx.py")
    findings = [
        checker.Finding(
            "WARN",
            "sections 12, 13, 14",
            "多个 section 的偶数页页眉在 python-docx 视图中重复为“附录”；需用 Word 或学校检测复核",
        ),
        checker.Finding("WARN", "section 12", "检测到横向 A4 section"),
    ]

    filtered = precheck.filter_findings_for_report(
        findings,
        {"available": True, "sections": []},
        [],
    )

    assert len(filtered) == 1
    assert filtered[0].message == "检测到横向 A4 section"
