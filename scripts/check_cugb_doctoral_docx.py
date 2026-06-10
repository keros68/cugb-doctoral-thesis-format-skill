#!/usr/bin/env python
"""Heuristic checker for CUGB doctoral thesis DOCX files."""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document


EXPECTED_PAGE = {
    "page_width": 21.0,
    "page_height": 29.7,
    "top_margin": 2.5,
    "bottom_margin": 2.0,
    "left_margin": 2.5,
    "right_margin": 2.0,
    "header_distance": 1.5,
    "footer_distance": 1.5,
}

CHINESE_RE = re.compile(r"[\u4e00-\u9fff]")
HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)(\S)")
CHAPTER_RE = re.compile(r"^(\d+)\s+(.+)")
FIGURE_RE = re.compile(r"^图\s*(\d+)[-－](\d+)")
TABLE_RE = re.compile(r"^表\s*(\d+)[-－](\d+)")
REFERENCE_RE = re.compile(r"^\[(\d+)\]")
COVER_FIELD_RE = re.compile(r"^(学\s*号|研\s*究\s*生|专\s*业|研\s*究\s*方\s*向|指\s*导\s*教\s*师|副\s*指\s*导\s*教\s*师)")


@dataclass
class Finding:
    level: str
    location: str
    message: str


def cm(value):
    return None if value is None else value.cm


def near(actual, expected, tolerance=0.08):
    return actual is not None and abs(actual - expected) <= tolerance


def clean(text: str) -> str:
    return " ".join(text.split())


def paragraph_text(paragraph) -> str:
    return clean(paragraph.text)


def emit(findings, level, location, message):
    findings.append(Finding(level, location, message))


def check_sections(doc, findings):
    attrs = {
        "page_width": "纸张宽度",
        "page_height": "纸张高度",
        "top_margin": "上边距",
        "bottom_margin": "下边距",
        "left_margin": "左边距",
        "right_margin": "右边距",
        "header_distance": "页眉距边界",
        "footer_distance": "页脚距边界",
    }
    even_headers = []
    for index, section in enumerate(doc.sections, 1):
        width = cm(section.page_width)
        height = cm(section.page_height)
        is_landscape_a4 = near(width, 29.7) and near(height, 21.0)
        if is_landscape_a4:
            emit(findings, "WARN", f"section {index}", "检测到横向 A4 section；若用于宽表/附录可保留，但需人工确认学校检测是否接受")
        for attr, label in attrs.items():
            actual = cm(getattr(section, attr))
            expected = EXPECTED_PAGE[attr]
            if is_landscape_a4 and attr in {"page_width", "page_height", "left_margin", "right_margin", "bottom_margin"}:
                continue
            if not near(actual, expected):
                emit(findings, "ERROR", f"section {index}", f"{label}应约为 {expected:.1f} cm，实际 {actual:.2f} cm")

        header = clean(" ".join(p.text for p in section.header.paragraphs))
        even_header = clean(" ".join(p.text for p in section.even_page_header.paragraphs))
        if even_header:
            even_headers.append((index, even_header))
        if index >= 2 and header and "中国地质大学" not in header and "博士学位论文" not in header:
            emit(findings, "WARN", f"section {index}", f"奇数页页眉未见学校博士论文固定文字：{header[:60]}")
        if even_header and ("硕士" in even_header or "中国民航大学" in even_header):
            emit(findings, "ERROR", f"section {index}", f"偶数页页眉疑似残留其他模板文字：{even_header[:60]}")

    repeated = {}
    for index, even_header in even_headers:
        repeated.setdefault(even_header, []).append(index)
    for even_header, sections in repeated.items():
        if len(sections) >= 3 and even_header not in {"中国地质大学（北京）博士学位论文"}:
            where = ", ".join(str(item) for item in sections)
            emit(findings, "WARN", f"sections {where}", f"多个 section 的偶数页页眉在 python-docx 视图中重复为“{even_header}”；需用 Word 或学校检测复核")


def run_sizes(paragraph):
    values = []
    for run in paragraph.runs:
        if run.text.strip() and run.font.size is not None:
            values.append(round(run.font.size.pt, 1))
    return values


def has_nonzero_first_line(paragraph, threshold_cm=0.08):
    indent = cm(paragraph.paragraph_format.first_line_indent)
    return indent is not None and abs(indent) > threshold_cm


def line_spacing_number(paragraph):
    value = paragraph.paragraph_format.line_spacing
    return value if isinstance(value, float) else None


def check_cover(paragraphs, findings):
    for index, paragraph in enumerate(paragraphs[:45], 1):
        text = paragraph_text(paragraph)
        if not text:
            continue
        spacing = line_spacing_number(paragraph)
        if text in {"中国地质大学（北京）"} and spacing and abs(spacing - 1.0) > 0.05:
            emit(findings, "ERROR", f"paragraph {index}", "封面学校名行应保持单倍行距")
        if re.match(r"^\d{4}年\d{1,2}月$", text) and spacing and abs(spacing - 1.0) > 0.05:
            emit(findings, "ERROR", f"paragraph {index}", "封面日期行应保持单倍行距")
        if COVER_FIELD_RE.match(text):
            indent = cm(paragraph.paragraph_format.first_line_indent)
            if indent is not None and abs(indent - 2.22) > 0.15:
                emit(findings, "ERROR", f"paragraph {index}", f"封面字段首行缩进应约为 2.22 cm，实际 {indent:.2f} cm")


def run_font_names(paragraph):
    names = []
    for run in paragraph.runs:
        if run.text.strip() and run.font.name:
            names.append(run.font.name)
    return names


def check_toc(paragraphs, findings):
    in_toc = False
    for index, paragraph in enumerate(paragraphs, 1):
        text = paragraph_text(paragraph)
        if text in {"目 录", "目录"}:
            in_toc = True
            continue
        if in_toc and text == "1 引言":
            break
        if not in_toc or not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        looks_like_toc_item = style_name.startswith("目录") or re.search(r"\s+\d+$", text)
        if not looks_like_toc_item:
            continue
        names = run_font_names(paragraph)
        style_font = paragraph.style.font.name if paragraph.style and paragraph.style.font else None
        if any("黑体" in name for name in names) or (style_font and "黑体" in style_font):
            emit(findings, "ERROR", f"paragraph {index}", "目录项不应使用黑体；应检查目录样式是否继承了正文标题格式")


def check_keywords(text, index, findings):
    if text.startswith("关键词") or "关键词：" in text:
        if "；" in text or ";" in text:
            emit(findings, "ERROR", f"paragraph {index}", "中文关键词分隔符应为中文逗号，不应使用分号")
        if "，" not in text and "," in text:
            emit(findings, "WARN", f"paragraph {index}", "中文关键词建议使用中文逗号")
    if re.match(r"(?i)^key\s*words", text) or re.match(r"(?i)^keywords", text):
        if not text.startswith("Key words:"):
            emit(findings, "ERROR", f"paragraph {index}", "英文关键词标记应为精确的 'Key words:'")
        after = text.split(":", 1)[-1] if ":" in text else text
        if ";" in after or "；" in after:
            emit(findings, "ERROR", f"paragraph {index}", "英文关键词应使用半角逗号分隔，不应使用分号")


def check_heading(text, paragraph, index, findings):
    if not looks_like_numbered_heading(text, paragraph):
        return
    match = HEADING_RE.match(text)
    if match and not re.match(r"^\d+(?:\.\d+)*\s+", text):
        emit(findings, "ERROR", f"paragraph {index}", "标题编号与标题文字之间应有 1 个空格")
    if re.match(r"^\d+(?:\.\d+)*\s+", text) and has_nonzero_first_line(paragraph):
        emit(findings, "ERROR", f"paragraph {index}", "标题段落不应保留正文首行缩进")


def looks_like_numbered_heading(text, paragraph):
    if not HEADING_RE.match(text) and not re.match(r"^\d+(?:\.\d+)*\s+", text):
        return False
    if re.match(r"^\d{4}年", text):
        return False
    if re.match(r"^\d+\.\s", text):
        return False
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading") or "标题" in style_name:
        return True
    sizes = run_sizes(paragraph)
    if sizes and max(sizes) >= 14:
        return True
    text_len = len(text)
    return text_len <= 45 and paragraph.paragraph_format.first_line_indent is None


def check_caption(text, paragraph, index, findings, captions):
    fig = FIGURE_RE.match(text)
    table = TABLE_RE.match(text)
    kind = None
    match = None
    if fig:
        kind, match = "图", fig
    elif table:
        kind, match = "表", table
    if not match:
        return

    chapter, number = int(match.group(1)), int(match.group(2))
    captions.setdefault(kind, []).append((chapter, number, index, text))
    if has_nonzero_first_line(paragraph):
        emit(findings, "ERROR", f"paragraph {index}", f"{kind}题不应有首行缩进")
    if text[-1:] in "，。；：,.、;:":
        emit(findings, "ERROR", f"paragraph {index}", f"{kind}题末尾通常不应加标点")
    sizes = run_sizes(paragraph)
    if sizes and any(abs(size - 10.5) > 0.2 for size in sizes):
        emit(findings, "WARN", f"paragraph {index}", f"{kind}题字号通常应为五号 10.5 pt，检测到 {sorted(set(sizes))}")


def check_chinese_punctuation(text, index, findings, in_references=False):
    if in_references or not CHINESE_RE.search(text):
        return
    if '"' in text:
        emit(findings, "REMIND", f"paragraph {index}", "中文段落中出现半角双引号，通常应改为中文引号")
    if ";" in text:
        emit(findings, "REMIND", f"paragraph {index}", "中文段落中出现半角分号，通常应改为中文分号")
    if re.search(r"[\u4e00-\u9fff],[\u4e00-\u9fff]", text):
        emit(findings, "REMIND", f"paragraph {index}", "中文词语之间出现半角逗号，通常应改为中文逗号")
    if "（" in text and ")" in text or "(" in text and "）" in text:
        emit(findings, "REMIND", f"paragraph {index}", "括号全角/半角混用")


def check_caption_sequence(captions, findings):
    for kind, rows in captions.items():
        by_chapter = {}
        for chapter, number, index, text in rows:
            by_chapter.setdefault(chapter, []).append((number, index, text))
        for chapter, items in by_chapter.items():
            seen = set()
            expected = 1
            for number, index, _ in sorted(items):
                if number in seen:
                    emit(findings, "ERROR", f"paragraph {index}", f"{kind}{chapter}-{number} 编号重复")
                seen.add(number)
                if number != expected:
                    emit(findings, "WARN", f"paragraph {index}", f"{kind}{chapter}-{number} 可能不连续，期望 {kind}{chapter}-{expected}")
                    expected = number
                expected += 1


def check_references(paragraphs, findings):
    start = None
    for i, paragraph in enumerate(paragraphs):
        if clean(paragraph.text) == "参考文献":
            start = i + 1
            break
    if start is None:
        emit(findings, "WARN", "document", "未找到独立的参考文献标题")
        return

    expected = 1
    for offset, paragraph in enumerate(paragraphs[start:], start + 1):
        text = paragraph_text(paragraph)
        if not text:
            continue
        if text in {"致谢", "附录", "个人简历"}:
            break
        match = REFERENCE_RE.match(text)
        if not match:
            continue
        number = int(match.group(1))
        if number != expected:
            emit(findings, "ERROR", f"paragraph {offset}", f"参考文献编号不连续，期望 [{expected}]，实际 [{number}]")
            expected = number
        expected += 1
        if "[J]" in text and "：" not in text and ":" not in text:
            emit(findings, "REMIND", f"paragraph {offset}", "期刊参考文献可能缺少年卷期后的页码")
        if re.search(r"[A-Za-z],\S", text):
            emit(findings, "REMIND", f"paragraph {offset}", "英文参考文献逗号后可能缺少空格")


def xml_counts(path: Path):
    counts = {"toc": 0, "page": 0, "omath": 0}
    with zipfile.ZipFile(path) as package:
        for name in package.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            data = package.read(name).decode("utf-8", errors="ignore")
            counts["toc"] += data.count("TOC ")
            counts["page"] += data.count(" PAGE ")
            counts["omath"] += data.count("<m:oMath")
    return counts


def check_xml(path, findings):
    counts = xml_counts(path)
    if counts["toc"] == 0:
        emit(findings, "WARN", "document.xml", "未检测到 Word 自动目录 TOC 域；若目录为手打文本，需人工逐项核对页码")
    if counts["page"] == 0:
        emit(findings, "WARN", "document.xml", "未检测到 PAGE 页码域；需确认页码是否为真实字段")
    emit(findings, "INFO", "document.xml", f"TOC fields={counts['toc']}, PAGE fields={counts['page']}, OMath objects={counts['omath']}")


def check_document(path: Path):
    doc = Document(path)
    findings = []
    captions = {}
    check_sections(doc, findings)
    check_cover(doc.paragraphs, findings)
    check_toc(doc.paragraphs, findings)

    in_references = False
    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph_text(paragraph)
        if not text:
            continue
        if text == "参考文献":
            in_references = True
        check_keywords(text, index, findings)
        check_heading(text, paragraph, index, findings)
        check_caption(text, paragraph, index, findings, captions)
        check_chinese_punctuation(text, index, findings, in_references)

    check_caption_sequence(captions, findings)
    check_references(doc.paragraphs, findings)
    check_xml(path, findings)
    return findings


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument("--fail-on-error", action="store_true")
    args = parser.parse_args()

    findings = check_document(args.docx)
    order = {"ERROR": 0, "WARN": 1, "REMIND": 2, "INFO": 3}
    findings.sort(key=lambda item: (order.get(item.level, 9), item.location, item.message))
    for item in findings:
        print(f"[{item.level}] {item.location}: {item.message}")

    error_count = sum(1 for item in findings if item.level == "ERROR")
    warn_count = sum(1 for item in findings if item.level == "WARN")
    remind_count = sum(1 for item in findings if item.level == "REMIND")
    print(f"Summary: errors={error_count}, warnings={warn_count}, reminders={remind_count}")
    if args.fail_on_error and error_count:
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
